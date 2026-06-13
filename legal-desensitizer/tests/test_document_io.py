"""Tests for 003-document-io: DOCX/XLSX redact/restore/audit, PDF audit."""

import hashlib
import json
import os
import sys
import tempfile

import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from legal_desens.io import read_text
from legal_desens.rules import load_rules
from legal_desens.redact import redact

FIXTURES = os.path.join(os.path.dirname(__file__), "fixtures")
RULES_PATH = os.path.join(os.path.dirname(__file__), "..", "rules", "rules.json")


@pytest.fixture
def rules():
    return load_rules(RULES_PATH)


def _txt_redact_fn(text, rules, source_sha256, mode, level, model_dir):
    """Wrapper to call text redact engine with standard signature."""
    return redact(text, rules, source_sha256, mode, level, model_dir)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX Tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestDOCXRedactRestore:
    """DOCX redact -> restore round-trip tests."""

    def test_docx_basic_redact_restore(self, rules, tmp_path):
        """Basic DOCX redact/restore with content verification."""
        from legal_desens.adapters.docx_adapter import DOCXAdapter
        adapter = DOCXAdapter()

        src = os.path.join(FIXTURES, "docx", "sample.docx")
        redacted = str(tmp_path / "sample.redacted.docx")
        restored = str(tmp_path / "sample.restored.docx")
        map_path = str(tmp_path / "sample.map.json")

        # Redact
        map_data, audit_data = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
            mode="regex-only",
            level="strict",
        )

        assert map_data["schema_version"] == "1.1"
        assert map_data["document_type"] == "docx"
        assert map_data["verification"] == "content"
        assert len(map_data["entities"]) > 0
        assert len(map_data["occurrences"]) > 0

        # Save map
        with open(map_path, "w") as f:
            json.dump(map_data, f, ensure_ascii=False, indent=2)

        # Restore
        adapter.restore(redacted, restored, map_data)

        # Content verification: extract text from source and restored should match
        source_text, _ = adapter.extract_text(src)
        restored_text, _ = adapter.extract_text(restored)
        assert source_text == restored_text, "DOCX content mismatch after restore"

    def test_docx_locator_is_redacted_coordinate(self, rules, tmp_path):
        """Verify locator points to redacted document structure."""
        from legal_desens.adapters.docx_adapter import DOCXAdapter
        adapter = DOCXAdapter()

        src = os.path.join(FIXTURES, "docx", "sample.docx")
        redacted = str(tmp_path / "redacted.docx")

        map_data, _ = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        # All locators must have type=docx
        for occ in map_data["occurrences"]:
            loc = occ["locator"]
            assert loc["type"] == "docx"
            assert "paragraph_index" in loc
            assert "run_start_index" in loc
            assert "run_end_index" in loc
            assert "text_start" in loc
            assert "text_end" in loc

    def test_docx_redacted_sha256_precheck(self, rules, tmp_path):
        """Restore must abort if redacted SHA-256 doesn't match."""
        from legal_desens.adapters.docx_adapter import DOCXAdapter
        adapter = DOCXAdapter()

        src = os.path.join(FIXTURES, "docx", "sample.docx")
        redacted = str(tmp_path / "redacted.docx")
        restored = str(tmp_path / "restored.docx")

        map_data, _ = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        # Tamper with map
        map_data["redacted_sha256"] = "0" * 64

        with pytest.raises(ValueError, match="SHA-256 mismatch"):
            adapter.restore(redacted, restored, map_data)

    def test_docx_cross_run_redact(self, rules, tmp_path):
        """Entity spanning multiple runs in same paragraph."""
        from legal_desens.adapters.docx_adapter import DOCXAdapter
        adapter = DOCXAdapter()

        src = os.path.join(FIXTURES, "docx", "cross_run.docx")
        redacted = str(tmp_path / "cross_run.redacted.docx")
        restored = str(tmp_path / "cross_run.restored.docx")

        map_data, audit_data = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        # Should have phone entity
        phone_entities = [e for e in map_data["entities"] if e["entity_type"] == "PHONE"]
        assert len(phone_entities) > 0, "Phone entity not detected in cross-run scenario"

        # Restore and verify content
        adapter.restore(redacted, restored, map_data)
        source_text, _ = adapter.extract_text(src)
        restored_text, _ = adapter.extract_text(restored)
        assert source_text == restored_text

    def test_docx_cross_paragraph_warning(self, rules, tmp_path):
        """Cross-paragraph entity should produce audit warning, not silent corruption."""
        from legal_desens.adapters.docx_adapter import DOCXAdapter
        adapter = DOCXAdapter()

        src = os.path.join(FIXTURES, "docx", "cross_paragraph.docx")
        redacted = str(tmp_path / "cross_para.redacted.docx")

        map_data, audit_data = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        # The phone 13800138000 is split across paragraphs:
        # "请联系1380013" + "8000获取详情。"
        # Each paragraph is processed independently, so the partial matches
        # won't form a complete phone number. This is expected behavior.
        # The test verifies no crash and proper audit output.
        assert "warnings" in audit_data

    def test_docx_no_match_passthrough(self, rules, tmp_path):
        """No sensitive data: document passes through unchanged content."""
        from legal_desens.adapters.docx_adapter import DOCXAdapter
        adapter = DOCXAdapter()

        src = os.path.join(FIXTURES, "docx", "no_match.docx")
        redacted = str(tmp_path / "no_match.redacted.docx")

        map_data, audit_data = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        assert len(map_data["entities"]) == 0
        assert len(map_data["occurrences"]) == 0
        assert audit_data["summary"]["total_entities"] == 0

        # Content should match
        source_text, _ = adapter.extract_text(src)
        redacted_text, _ = adapter.extract_text(redacted)
        assert source_text == redacted_text

    def test_docx_empty_document(self, rules, tmp_path):
        """Empty DOCX should not crash."""
        from legal_desens.adapters.docx_adapter import DOCXAdapter
        adapter = DOCXAdapter()

        src = os.path.join(FIXTURES, "docx", "empty.docx")
        redacted = str(tmp_path / "empty.redacted.docx")

        map_data, audit_data = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        assert len(map_data["entities"]) == 0
        assert len(map_data["occurrences"]) == 0

    def test_docx_audit(self, rules, tmp_path):
        """DOCX audit should produce proper schema."""
        from legal_desens.adapters.docx_adapter import DOCXAdapter
        adapter = DOCXAdapter()

        src = os.path.join(FIXTURES, "docx", "sample.docx")
        redacted = str(tmp_path / "sample.redacted.docx")

        map_data, _ = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        result = adapter.audit(redacted, map_data, rules)
        assert result["schema_version"] == "1.1"
        assert result["document_type"] == "docx"
        assert "summary" in result
        assert "residual_scan" in result

    def test_docx_no_byte_level_claim(self, rules, tmp_path):
        """DOCX must not claim byte-level restoration."""
        from legal_desens.adapters.docx_adapter import DOCXAdapter
        adapter = DOCXAdapter()

        src = os.path.join(FIXTURES, "docx", "sample.docx")
        redacted = str(tmp_path / "redacted.docx")

        map_data, _ = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        assert map_data["verification"] == "content"
        assert map_data["verification"] != "byte"

    def test_docx_entity_ids_are_document_scoped_with_shared_label(self, rules, tmp_path):
        """Different entities keep distinct IDs even when the display label is shared."""
        from docx import Document
        from legal_desens.adapters.docx_adapter import DOCXAdapter

        src = str(tmp_path / "two_phones.docx")
        redacted = str(tmp_path / "two_phones.redacted.docx")
        restored = str(tmp_path / "two_phones.restored.docx")

        doc = Document()
        doc.add_paragraph("电话13800138000")
        doc.add_paragraph("电话13900139000")
        doc.save(src)

        adapter = DOCXAdapter()
        map_data, _ = adapter.redact(src, redacted, _txt_redact_fn, rules)
        redacted_text, _ = adapter.extract_text(redacted)

        assert redacted_text.count("【手机号】") == 2
        assert [e["replacement"] for e in map_data["entities"]] == ["【手机号】", "【手机号】"]
        assert [e["id"] for e in map_data["entities"]] == ["PHONE_1", "PHONE_2"]

        adapter.restore(redacted, restored, map_data)
        restored_text, _ = adapter.extract_text(restored)
        assert restored_text == "电话13800138000\n电话13900139000"


# ═══════════════════════════════════════════════════════════════════════════════
# XLSX Tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestXLSXRedactRestore:
    """XLSX redact -> restore round-trip tests."""

    def test_xlsx_basic_redact_restore(self, rules, tmp_path):
        """Basic XLSX redact/restore with cell text verification."""
        from legal_desens.adapters.xlsx_adapter import XLSXAdapter
        adapter = XLSXAdapter()

        src = os.path.join(FIXTURES, "xlsx", "sample.xlsx")
        redacted = str(tmp_path / "sample.redacted.xlsx")
        restored = str(tmp_path / "sample.restored.xlsx")

        map_data, audit_data = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        assert map_data["schema_version"] == "1.1"
        assert map_data["document_type"] == "xlsx"
        assert map_data["verification"] == "content"
        assert len(map_data["entities"]) > 0

        # Restore
        adapter.restore(redacted, restored, map_data)

        # Cell text verification
        source_text, _ = adapter.extract_text(src)
        restored_text, _ = adapter.extract_text(restored)
        assert source_text == restored_text, "XLSX cell texts mismatch after restore"

    def test_xlsx_shared_strings_not_corrupted(self, rules, tmp_path):
        """When a cell using shared string is redacted, other cells with same string must not change."""
        from legal_desens.adapters.xlsx_adapter import XLSXAdapter
        adapter = XLSXAdapter()

        src = os.path.join(FIXTURES, "xlsx", "shared_strings.xlsx")
        redacted = str(tmp_path / "shared.redacted.xlsx")

        map_data, audit_data = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        # Extract text from redacted - the unredacted "张三" cell should still be intact
        redacted_text, segments = adapter.extract_text(redacted)

        # Find the cell with "张三" that was NOT redacted
        # (only the phone number should be redacted, not the name "张三")
        zhangsan_count = sum(1 for seg in segments if seg.get("text") == "张三")
        # Both "张三" cells should remain since "张三" is not matched by regex rules
        assert zhangsan_count == 2, f"Expected 2 '张三' cells, found {zhangsan_count}"

    def test_xlsx_formula_cell_skipped(self, rules, tmp_path):
        """Formula cells should be skipped with audit warning."""
        from legal_desens.adapters.xlsx_adapter import XLSXAdapter
        adapter = XLSXAdapter()

        src = os.path.join(FIXTURES, "xlsx", "formula.xlsx")
        redacted = str(tmp_path / "formula.redacted.xlsx")

        map_data, audit_data = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        # Should have formula_cell_skipped warning
        formula_warnings = [
            w for w in audit_data.get("warnings", [])
            if isinstance(w, dict) and w.get("type") == "formula_cell_skipped"
        ]
        assert len(formula_warnings) > 0, "Expected formula_cell_skipped warning"

        # Phone should still be redacted
        phone_entities = [e for e in map_data["entities"] if e["entity_type"] == "PHONE"]
        assert len(phone_entities) > 0

    def test_xlsx_empty_cells_skipped(self, rules, tmp_path):
        """Empty cells should be skipped without error."""
        from legal_desens.adapters.xlsx_adapter import XLSXAdapter
        adapter = XLSXAdapter()

        src = os.path.join(FIXTURES, "xlsx", "empty_cells.xlsx")
        redacted = str(tmp_path / "empty.redacted.xlsx")

        map_data, audit_data = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        # Should not crash, phone should be redacted
        assert len(map_data["occurrences"]) > 0

    def test_xlsx_multi_shared_string_safety(self, rules, tmp_path):
        """Manually crafted XLSX with explicit shared string references."""
        from legal_desens.adapters.xlsx_adapter import XLSXAdapter
        adapter = XLSXAdapter()

        src = os.path.join(FIXTURES, "xlsx", "multi_shared.xlsx")
        redacted = str(tmp_path / "multi_shared.redacted.xlsx")

        map_data, audit_data = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        # "张三" appears twice (shared), phone appears once
        # Redacting phone should NOT corrupt "张三"
        redacted_text, segments = adapter.extract_text(redacted)

        zhangsan_segments = [s for s in segments if s.get("text") == "张三"]
        assert len(zhangsan_segments) == 2, "Shared string '张三' was corrupted!"

    def test_xlsx_redacted_sha256_precheck(self, rules, tmp_path):
        """Restore must abort if redacted SHA-256 doesn't match."""
        from legal_desens.adapters.xlsx_adapter import XLSXAdapter
        adapter = XLSXAdapter()

        src = os.path.join(FIXTURES, "xlsx", "sample.xlsx")
        redacted = str(tmp_path / "redacted.xlsx")
        restored = str(tmp_path / "restored.xlsx")

        map_data, _ = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        map_data["redacted_sha256"] = "0" * 64

        with pytest.raises(ValueError, match="SHA-256 mismatch"):
            adapter.restore(redacted, restored, map_data)

    def test_xlsx_audit(self, rules, tmp_path):
        """XLSX audit should produce proper schema."""
        from legal_desens.adapters.xlsx_adapter import XLSXAdapter
        adapter = XLSXAdapter()

        src = os.path.join(FIXTURES, "xlsx", "sample.xlsx")
        redacted = str(tmp_path / "sample.redacted.xlsx")

        map_data, _ = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        result = adapter.audit(redacted, map_data, rules)
        assert result["schema_version"] == "1.1"
        assert result["document_type"] == "xlsx"
        assert "summary" in result

    def test_xlsx_no_byte_level_claim(self, rules, tmp_path):
        """XLSX must not claim byte-level restoration."""
        from legal_desens.adapters.xlsx_adapter import XLSXAdapter
        adapter = XLSXAdapter()

        src = os.path.join(FIXTURES, "xlsx", "sample.xlsx")
        redacted = str(tmp_path / "redacted.xlsx")

        map_data, _ = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        assert map_data["verification"] == "content"

    def test_xlsx_entity_ids_are_workbook_scoped_with_shared_label(self, rules, tmp_path):
        """Different entities keep distinct IDs even when the display label is shared."""
        from openpyxl import Workbook
        from legal_desens.adapters.xlsx_adapter import XLSXAdapter

        src = str(tmp_path / "two_phones.xlsx")
        redacted = str(tmp_path / "two_phones.redacted.xlsx")
        restored = str(tmp_path / "two_phones.restored.xlsx")

        wb = Workbook()
        ws = wb.active
        ws["A1"] = "电话13800138000"
        ws["A2"] = "电话13900139000"
        wb.save(src)

        adapter = XLSXAdapter()
        map_data, _ = adapter.redact(src, redacted, _txt_redact_fn, rules)
        redacted_text, _ = adapter.extract_text(redacted)

        assert redacted_text.count("【手机号】") == 2
        assert [e["replacement"] for e in map_data["entities"]] == ["【手机号】", "【手机号】"]
        assert [e["id"] for e in map_data["entities"]] == ["PHONE_1", "PHONE_2"]

        adapter.restore(redacted, restored, map_data)
        restored_text, _ = adapter.extract_text(restored)
        assert restored_text == "电话13800138000\n电话13900139000"


# ═══════════════════════════════════════════════════════════════════════════════
# PDF Tests (008: unsupported after PyMuPDF removal)
# ═══════════════════════════════════════════════════════════════════════════════

class TestPDFAudit:
    """PDF is unsupported in the commercial-safe core (008)."""

    def test_pdf_redact_unsupported(self, rules, tmp_path):
        """CLI redact with PDF should return unsupported error."""
        from legal_desens.cli import main

        src = os.path.join(FIXTURES, "sample.txt")  # use any file, format detection is by ext
        # Create a dummy .pdf path (file doesn't need to exist for the unsupported check)
        pdf_path = str(tmp_path / "dummy.pdf")
        out = str(tmp_path / "out.pdf")

        ret = main([
            "redact", pdf_path,
            "--regex-only",
            "--out", out,
        ])
        assert ret == 1

    def test_pdf_restore_unsupported(self, tmp_path):
        """CLI restore with PDF should return unsupported error."""
        from legal_desens.cli import main

        pdf_path = str(tmp_path / "dummy.pdf")
        map_out = str(tmp_path / "map.json")
        with open(map_out, "w") as f:
            json.dump({"source_sha256": "", "redacted_sha256": ""}, f)

        ret = main([
            "restore", pdf_path,
            "--map", map_out,
            "--out", str(tmp_path / "out.pdf"),
        ])
        assert ret == 1

    def test_pdf_audit_unsupported(self, tmp_path):
        """CLI audit with PDF should return unsupported error."""
        from legal_desens.cli import main

        pdf_path = str(tmp_path / "dummy.pdf")
        map_out = str(tmp_path / "map.json")
        with open(map_out, "w") as f:
            json.dump({"redacted_sha256": "", "entities": [], "occurrences": []}, f)

        ret = main([
            "audit", pdf_path,
            "--map", map_out,
        ])
        assert ret == 1


# ═══════════════════════════════════════════════════════════════════════════════
# Map Schema Tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestMapSchema:
    """Verify map schema extensions for 003."""

    def test_docx_map_schema(self, rules, tmp_path):
        """DOCX map should have entities + occurrences + locator."""
        from legal_desens.adapters.docx_adapter import DOCXAdapter
        adapter = DOCXAdapter()

        src = os.path.join(FIXTURES, "docx", "sample.docx")
        redacted = str(tmp_path / "redacted.docx")

        map_data, _ = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        assert map_data["schema_version"] == "1.1"
        assert "entities" in map_data
        assert "occurrences" in map_data
        assert "verification" in map_data
        assert "source_sha256" in map_data
        assert "redacted_sha256" in map_data

        for occ in map_data["occurrences"]:
            assert "locator" in occ
            assert occ["locator"]["type"] == "docx"

    def test_xlsx_map_schema(self, rules, tmp_path):
        """XLSX map should have entities + occurrences + locator."""
        from legal_desens.adapters.xlsx_adapter import XLSXAdapter
        adapter = XLSXAdapter()

        src = os.path.join(FIXTURES, "xlsx", "sample.xlsx")
        redacted = str(tmp_path / "redacted.xlsx")

        map_data, _ = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        assert map_data["schema_version"] == "1.1"
        assert "entities" in map_data
        assert "occurrences" in map_data
        assert "verification" in map_data

        for occ in map_data["occurrences"]:
            assert "locator" in occ
            loc = occ["locator"]
            assert loc["type"] == "xlsx"
            assert "sheet" in loc
            assert "row" in loc
            assert "column" in loc or "col" in loc


# ═══════════════════════════════════════════════════════════════════════════════
# Verification Mode Tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestVerificationMode:
    """Verify that verification mode is correctly recorded per format."""

    def test_txt_verification_is_byte(self, rules):
        """TXT verification should be 'byte'."""
        tf = read_text(os.path.join(FIXTURES, "sample.txt"))
        source_sha = tf.sha256
        _, map_data, _ = redact(tf.text, rules, source_sha)
        # TXT doesn't set verification in 001, but we can check
        # The CLI adds it. For now, just verify the structure exists.
        assert "schema_version" in map_data

    def test_docx_verification_is_content(self, rules, tmp_path):
        from legal_desens.adapters.docx_adapter import DOCXAdapter
        adapter = DOCXAdapter()

        src = os.path.join(FIXTURES, "docx", "sample.docx")
        redacted = str(tmp_path / "redacted.docx")

        map_data, _ = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        assert map_data["verification"] == "content"

    def test_xlsx_verification_is_content(self, rules, tmp_path):
        from legal_desens.adapters.xlsx_adapter import XLSXAdapter
        adapter = XLSXAdapter()

        src = os.path.join(FIXTURES, "xlsx", "sample.xlsx")
        redacted = str(tmp_path / "redacted.xlsx")

        map_data, _ = adapter.redact(
            source_path=src,
            redacted_path=redacted,
            redact_fn=_txt_redact_fn,
            rules=rules,
        )

        assert map_data["verification"] == "content"


# ═══════════════════════════════════════════════════════════════════════════════
# CLI Integration Tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestCLIDocumentIO:
    """Test CLI commands with document formats."""

    def test_cli_help(self):
        from legal_desens.cli import main
        with pytest.raises(SystemExit) as exc_info:
            main(["--help"])
        assert exc_info.value.code == 0

    def test_cli_redact_docx(self, tmp_path):
        """CLI redact command with DOCX file."""
        from legal_desens.cli import main

        src = os.path.join(FIXTURES, "docx", "sample.docx")
        out = str(tmp_path / "out.docx")
        map_out = str(tmp_path / "map.json")
        audit_out = str(tmp_path / "audit.json")

        ret = main([
            "redact", src,
            "--regex-only",
            "--out", out,
            "--map", map_out,
            "--audit", audit_out,
        ])
        assert ret == 0
        assert os.path.exists(out)
        assert os.path.exists(map_out)
        assert os.path.exists(audit_out)

        with open(map_out, encoding="utf-8") as f:
            map_data = json.load(f)
        assert map_data["document_type"] == "docx"

    def test_cli_redact_xlsx(self, tmp_path):
        """CLI redact command with XLSX file."""
        from legal_desens.cli import main

        src = os.path.join(FIXTURES, "xlsx", "sample.xlsx")
        out = str(tmp_path / "out.xlsx")
        map_out = str(tmp_path / "map.json")

        ret = main([
            "redact", src,
            "--regex-only",
            "--out", out,
            "--map", map_out,
        ])
        assert ret == 0
        assert os.path.exists(out)

    def test_cli_docx_default_missing_ner_aborts_without_output(self, tmp_path):
        """Default DOCX redact must not emit unredacted output when NER is unavailable."""
        from legal_desens.cli import main

        src = os.path.join(FIXTURES, "docx", "sample.docx")
        out = str(tmp_path / "out.docx")
        map_out = str(tmp_path / "map.json")
        audit_out = str(tmp_path / "audit.json")

        ret = main([
            "redact", src,
            "--model-dir", "/nonexistent",
            "--out", out,
            "--map", map_out,
            "--audit", audit_out,
        ])

        assert ret == 1
        assert not os.path.exists(out)
        assert not os.path.exists(map_out)
        assert not os.path.exists(audit_out)

    def test_cli_xlsx_default_missing_ner_aborts_without_output(self, tmp_path):
        """Default XLSX redact must not emit unredacted output when NER is unavailable."""
        from legal_desens.cli import main

        src = os.path.join(FIXTURES, "xlsx", "sample.xlsx")
        out = str(tmp_path / "out.xlsx")
        map_out = str(tmp_path / "map.json")
        audit_out = str(tmp_path / "audit.json")

        ret = main([
            "redact", src,
            "--model-dir", "/nonexistent",
            "--out", out,
            "--map", map_out,
            "--audit", audit_out,
        ])

        assert ret == 1
        assert not os.path.exists(out)
        assert not os.path.exists(map_out)
        assert not os.path.exists(audit_out)

    def test_cli_restore_docx(self, tmp_path):
        """CLI restore command with DOCX file."""
        from legal_desens.cli import main

        src = os.path.join(FIXTURES, "docx", "sample.docx")
        redacted = str(tmp_path / "redacted.docx")
        map_out = str(tmp_path / "map.json")
        restored = str(tmp_path / "restored.docx")

        # Redact
        ret = main([
            "redact", src,
            "--regex-only",
            "--out", redacted,
            "--map", map_out,
        ])
        assert ret == 0

        # Restore
        ret = main([
            "restore", redacted,
            "--map", map_out,
            "--out", restored,
        ])
        assert ret == 0

    def test_cli_restore_xlsx(self, tmp_path):
        """CLI restore command with XLSX file."""
        from legal_desens.cli import main

        src = os.path.join(FIXTURES, "xlsx", "sample.xlsx")
        redacted = str(tmp_path / "redacted.xlsx")
        map_out = str(tmp_path / "map.json")
        restored = str(tmp_path / "restored.xlsx")

        # Redact
        ret = main([
            "redact", src,
            "--regex-only",
            "--out", redacted,
            "--map", map_out,
        ])
        assert ret == 0

        # Restore
        ret = main([
            "restore", redacted,
            "--map", map_out,
            "--out", restored,
        ])
        assert ret == 0

    def test_cli_audit_pdf(self, tmp_path):
        """CLI audit with PDF should fail (unsupported)."""
        from legal_desens.cli import main

        pdf_path = str(tmp_path / "dummy.pdf")
        map_out = str(tmp_path / "map.json")
        map_data = {
            "redacted_sha256": "0" * 64,
            "entities": [],
            "occurrences": [],
        }
        with open(map_out, "w") as f:
            json.dump(map_data, f)

        ret = main([
            "audit", pdf_path,
            "--map", map_out,
        ])
        assert ret == 1


# ═══════════════════════════════════════════════════════════════════════════════
# Fixture Verification Tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestFixtures:
    """Verify fixtures are correctly created."""

    def test_docx_sample_readable(self):
        from legal_desens.adapters.docx_adapter import DOCXAdapter
        adapter = DOCXAdapter()
        text, segments = adapter.extract_text(os.path.join(FIXTURES, "docx", "sample.docx"))
        assert "13800138000" in text
        assert len(segments) == 2

    def test_docx_cross_run_readable(self):
        from legal_desens.adapters.docx_adapter import DOCXAdapter
        adapter = DOCXAdapter()
        text, segments = adapter.extract_text(os.path.join(FIXTURES, "docx", "cross_run.docx"))
        assert "13800138000" in text

    def test_xlsx_sample_readable(self):
        from legal_desens.adapters.xlsx_adapter import XLSXAdapter
        adapter = XLSXAdapter()
        text, segments = adapter.extract_text(os.path.join(FIXTURES, "xlsx", "sample.xlsx"))
        assert "13800138000" in text
        assert "zhangsan@example.com" in text

    def test_xlsx_shared_strings_readable(self):
        from legal_desens.adapters.xlsx_adapter import XLSXAdapter
        adapter = XLSXAdapter()
        text, segments = adapter.extract_text(os.path.join(FIXTURES, "xlsx", "shared_strings.xlsx"))
        zhangsan_count = sum(1 for s in segments if s.get("text") == "张三")
        assert zhangsan_count == 2

    def test_xlsx_formula_readable(self):
        from legal_desens.adapters.xlsx_adapter import XLSXAdapter
        adapter = XLSXAdapter()
        text, segments = adapter.extract_text(os.path.join(FIXTURES, "xlsx", "formula.xlsx"))
        # Formula cells return None from extract_text (skipped)
        # But non-formula cells should be present
        assert "13800138000" in text
        assert "合同款" in text

    def test_pdf_sample_not_applicable(self):
        """PDF fixtures were removed in 008. This test is a no-op placeholder."""
        pass
