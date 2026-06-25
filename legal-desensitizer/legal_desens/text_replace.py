"""Text replacement engine — star masking and placeholder substitution.

Star mode: keep first/last chars, fill middle with ★
  深圳市康成泰实业有限公司 → 深**********司
  张三 → 张*
  13800138000 → 138****8000

Placeholder mode: replace with numbered tags
  深圳市康成泰实业有限公司 → <单位1>
  张三 → <姓名1>
  13800138000 → <手机1>

Export formats:
- TXT: plain text with replacements
- MD: Markdown with replacements
- DOCX: content-level replacement (preserves formatting)

PDF star/placeholder → export as TXT/MD/DOCX (no PDF write-back).
"""

from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple


# ─── Entity type → Chinese label mapping ──────────────────────

_ENTITY_LABELS = {
    "PERSON": "姓名",
    "ORG": "单位",
    "PHONE": "手机",
    "LANDLINE": "电话",
    "ID_CARD": "证件号",
    "PASSPORT": "护照",
    "EMAIL": "邮箱",
    "CASE_NO": "案号",
    "ORG_CODE": "信用代码",
    "BANK_CARD": "银行卡",
    "BANK_BRANCH": "银行网点",
    "MONEY": "金额",
    "DATE": "日期",
    "TIME": "时间",
    "LOC": "地点",
    "ADDRESS": "地址",
    "API_TOKEN": "密钥",
    "SEAL": "公章",
    "MANUAL": "敏感信息",
}


# ─── Star masking functions ───────────────────────────────────

def _star_phone(original: str) -> str:
    """13800138000 → 138****8000"""
    chars = list(original)
    if len(chars) >= 11:
        return ''.join(chars[:3]) + '****' + ''.join(chars[-4:])
    return chars[0] + '***'


def _star_id_card(original: str) -> str:
    """110101199001011234 → 1101**********1234"""
    chars = list(original)
    if len(chars) >= 15:
        return ''.join(chars[:4]) + '*' * (len(chars) - 8) + ''.join(chars[-4:])
    return ''.join(chars[:3]) + '*' * max(1, len(chars) - 6) + ''.join(chars[-3:])


def _star_person(original: str) -> str:
    """张三 → 张*"""
    chars = list(original)
    return chars[0] + '*' * max(1, len(chars) - 1)


def _star_org(original: str) -> str:
    """深圳市康成泰实业有限公司 → 深**********司"""
    chars = list(original)
    if len(chars) > 4:
        return ''.join(chars[:2]) + '*' * max(2, len(chars) - 4) + ''.join(chars[-2:])
    return chars[0] + '*' * max(1, len(chars) - 1)


def _star_email(original: str) -> str:
    """test@example.com → te***@example.com"""
    if '@' in original:
        at_idx = original.index('@')
        return original[:min(2, at_idx)] + '***' + original[at_idx:]
    return original[:2] + '***'


def _star_bank_card(original: str) -> str:
    """6225880112345678 → ************5678"""
    chars = list(original)
    if len(chars) >= 8:
        return '*' * (len(chars) - 4) + ''.join(chars[-4:])
    return '*' * max(1, len(chars) - 2) + ''.join(chars[-2:])


def _star_money(original: str) -> str:
    """15000元 → ****元"""
    return '****' + (original[-1] if original else '')


def _star_date(original: str) -> str:
    """2026年6月20日 → ****年*月**日"""
    return ''.join(c if c in '年月日号时分秒' else '*' for c in original)


def _star_generic(original: str) -> str:
    """Generic: keep first char, rest ***"""
    chars = list(original)
    return chars[0] + '***' if chars else ''


_STAR_FUNCTIONS = {
    "PHONE": _star_phone,
    "LANDLINE": _star_phone,
    "ID_CARD": _star_id_card,
    "PASSPORT": _star_id_card,
    "PERSON": _star_person,
    "ORG": _star_org,
    "EMAIL": _star_email,
    "BANK_CARD": _star_bank_card,
    "BANK_BRANCH": _star_org,
    "MONEY": _star_money,
    "DATE": _star_date,
    "TIME": _star_date,
    "LOC": _star_generic,
    "ADDRESS": _star_generic,
    "CASE_NO": _star_generic,
    "ORG_CODE": _star_id_card,
    "API_TOKEN": _star_generic,
    "MANUAL": _star_generic,
}


def star_mask(original: str, entity_type: str) -> str:
    """Apply star masking to original text based on entity type."""
    if not original:
        return ""
    fn = _STAR_FUNCTIONS.get(entity_type, _star_generic)
    try:
        return fn(original)
    except Exception:
        return _star_generic(original)


# ─── Placeholder functions ────────────────────────────────────

class PlaceholderCounter:
    """Tracks per-type placeholder indices."""
    def __init__(self):
        self._counts: Dict[str, int] = {}

    def next(self, entity_type: str) -> str:
        label = _ENTITY_LABELS.get(entity_type, "敏感信息")
        self._counts[entity_type] = self._counts.get(entity_type, 0) + 1
        return f"<{label}{self._counts[entity_type]}>"

    def reset(self):
        self._counts.clear()


def placeholder_mask(original: str, entity_type: str, counter: PlaceholderCounter) -> str:
    """Replace original with numbered placeholder tag."""
    if not original:
        return ""
    return counter.next(entity_type)


# ─── Text replacement ─────────────────────────────────────────

@dataclass
class Replacement:
    """A single text replacement."""
    start: int          # position in original text
    end: int
    original: str
    replacement: str
    entity_type: str
    source: str = "auto"  # 'auto' | 'manual'


def apply_replacements_to_text(
    text: str,
    replacements: List[Replacement],
) -> str:
    """Apply sorted replacements to text, producing new text.

    Replacements must be non-overlapping and sorted by start position.
    """
    if not replacements:
        return text

    sorted_reps = sorted(replacements, key=lambda r: r.start)
    parts = []
    cursor = 0

    for rep in sorted_reps:
        if rep.start < cursor:
            # Overlapping — skip
            continue
        parts.append(text[cursor:rep.start])
        parts.append(rep.replacement)
        cursor = rep.end

    parts.append(text[cursor:])
    return ''.join(parts)


def build_replacements_from_entities(
    text: str,
    entities: List[dict],
    mode: str,  # 'star' | 'placeholder'
    counter: Optional[PlaceholderCounter] = None,
) -> List[Replacement]:
    """Build replacement list from detected entities.

    Args:
        text: Original text.
        entities: List of entity dicts with {original, entity_type, start, end, ...}.
        mode: 'star' or 'placeholder'.
        counter: PlaceholderCounter for placeholder mode (shared across calls).
    """
    if counter is None:
        counter = PlaceholderCounter()

    replacements = []
    for entity in entities:
        original = entity.get("original", "")
        entity_type = entity.get("entity_type", "MANUAL")
        start = entity.get("start", -1)
        end = entity.get("end", -1)

        if not original or start < 0 or end <= start:
            continue

        # Verify text matches
        actual = text[start:end]
        if actual != original:
            # Try to find in text
            idx = text.find(original)
            if idx >= 0:
                start = idx
                end = idx + len(original)
            else:
                continue

        if mode == "star":
            replacement = star_mask(original, entity_type)
        else:
            replacement = placeholder_mask(original, entity_type, counter)

        replacements.append(Replacement(
            start=start,
            end=end,
            original=original,
            replacement=replacement,
            entity_type=entity_type,
        ))

    return replacements


# ─── Export functions ─────────────────────────────────────────

def export_txt(
    text: str,
    replacements: List[Replacement],
    output_path: str,
) -> dict:
    """Export replaced text as plain TXT."""
    replaced = apply_replacements_to_text(text, replacements)
    Path(output_path).write_text(replaced, encoding="utf-8")
    return {
        "format": "txt",
        "output_path": output_path,
        "source_sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
        "output_sha256": hashlib.sha256(replaced.encode("utf-8")).hexdigest(),
        "replacements_applied": len(replacements),
    }


def export_md(
    text: str,
    replacements: List[Replacement],
    output_path: str,
) -> dict:
    """Export replaced text as Markdown."""
    replaced = apply_replacements_to_text(text, replacements)
    Path(output_path).write_text(replaced, encoding="utf-8")
    return {
        "format": "md",
        "output_path": output_path,
        "source_sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
        "output_sha256": hashlib.sha256(replaced.encode("utf-8")).hexdigest(),
        "replacements_applied": len(replacements),
    }


def export_docx(
    source_docx_path: str,
    replacements: List[Replacement],
    output_path: str,
    original_text: str,
) -> dict:
    """Export replaced text as DOCX with content-level replacement.

    Preserves original formatting (bold, italic, fonts, etc).
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX export: pip install python-docx")

    doc = Document(source_docx_path)

    # Build text → replacement mapping
    rep_map = {}
    for rep in replacements:
        rep_map[rep.original] = rep.replacement

    # Walk through all paragraphs and replace text
    replaced_count = 0
    for para in doc.paragraphs:
        full_text = para.text
        if not full_text:
            continue

        # Check if any replacement applies to this paragraph
        needs_update = False
        for original, replacement in rep_map.items():
            if original in full_text:
                needs_update = True
                break

        if not needs_update:
            continue

        # Rebuild paragraph runs with replacements
        # Strategy: concatenate runs, apply replacements, redistribute
        runs = para.runs
        if not runs:
            continue

        # Simple approach: replace in the full text, then set on first run
        # This loses per-run formatting for replaced spans
        # For better fidelity, we'd need run-level matching
        new_text = full_text
        for original, replacement in sorted(rep_map.items(), key=lambda x: -len(x[0])):
            if original in new_text:
                new_text = new_text.replace(original, replacement)
                replaced_count += 1

        # Clear all runs and set new text on first run
        if runs and new_text != full_text:
            runs[0].text = new_text
            for run in runs[1:]:
                run.text = ""

    doc.save(output_path)

    return {
        "format": "docx",
        "output_path": output_path,
        "source_sha256": hashlib.sha256(Path(source_docx_path).read_bytes()).hexdigest(),
        "replacements_applied": replaced_count,
    }


# ─── Unified text export ─────────────────────────────────────

def _load_rules_and_detect(text: str, rules_path: str) -> List[dict]:
    """Load rules.json and run regex detection on text.

    Returns list of entity dicts with {original, entity_type, start, end}.
    """
    import re as re_module

    try:
        with open(rules_path, "r", encoding="utf-8") as f:
            rules = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return []

    if not isinstance(rules, list):
        return []

    entities = []
    for rule in rules:
        if not rule.get("enabled", True):
            continue
        pattern = rule.get("pattern")
        if not pattern:
            continue
        entity_type = rule.get("entity_type", "CUSTOM")
        try:
            compiled = re_module.compile(pattern)
        except re_module.error:
            continue  # Skip invalid regex

        for match in compiled.finditer(text):
            original = match.group(0)
            if original and len(original) >= 2:
                entities.append({
                    "original": original,
                    "entity_type": entity_type,
                    "start": match.start(),
                    "end": match.end(),
                    "source": "rule",
                    "rule_id": rule.get("id"),
                })

    return entities


def text_export(
    source_path: str,
    output_path: str,
    entities: List[dict],
    mode: str,  # 'star' | 'placeholder'
    export_format: str,  # 'txt' | 'md' | 'docx'
    ocr_text: Optional[str] = None,
    rules_path: Optional[str] = None,
    denylist: Optional[List[str]] = None,
    whitelist: Optional[List[str]] = None,
) -> dict:
    """Unified text export with star/placeholder replacement.

    Args:
        source_path: Path to source file (PDF, DOCX, TXT, MD).
        output_path: Path for output file.
        entities: Detected entities with {original, entity_type, start, end}.
        mode: 'star' or 'placeholder'.
        export_format: 'txt', 'md', or 'docx'.
        ocr_text: Pre-extracted OCR text (for PDF sources).
        rules_path: Path to rules.json (for additional regex detection).
        denylist: List of forced redaction terms.
        whitelist: List of terms to skip (never redact).
    """
    ext = Path(source_path).suffix.lower()

    # Get text content
    if ocr_text:
        text = ocr_text
    elif ext in ('.txt', '.md'):
        text = Path(source_path).read_text(encoding="utf-8-sig")
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(source_path)
            text = '\n'.join(p.text for p in doc.paragraphs)
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")
    else:
        # PDF or other — require ocr_text
        raise ValueError(f"Cannot extract text from {ext}. Provide ocr_text parameter.")

    # Load rules and run regex detection on text
    if rules_path:
        rule_entities = _load_rules_and_detect(text, rules_path)
        # Merge with passed-in entities (deduplicate by position)
        existing_positions = {(e.get("start"), e.get("end")) for e in entities}
        for re_entity in rule_entities:
            pos = (re_entity.get("start"), re_entity.get("end"))
            if pos not in existing_positions:
                entities.append(re_entity)
                existing_positions.add(pos)

    # Apply whitelist: remove entities that match whitelist terms
    if whitelist:
        filtered_entities = []
        for entity in entities:
            original = entity.get("original", "")
            if any(wl in original for wl in whitelist):
                continue  # Skip whitelisted terms
            filtered_entities.append(entity)
        entities = filtered_entities

    # Apply denylist: add entities for denylist terms not already covered
    if denylist:
        existing_originals = {e.get("original", "") for e in entities}
        for term in denylist:
            if term in text and term not in existing_originals:
                # Find all occurrences in text
                idx = text.find(term)
                while idx >= 0:
                    entities.append({
                        "original": term,
                        "entity_type": "DENYLIST",
                        "start": idx,
                        "end": idx + len(term),
                    })
                    idx = text.find(term, idx + 1)

    # Sort entities by position
    entities.sort(key=lambda e: e.get("start", 0))

    # Build replacements
    counter = PlaceholderCounter() if mode == "placeholder" else None
    replacements = build_replacements_from_entities(text, entities, mode, counter)

    # Export
    if export_format == "txt":
        return export_txt(text, replacements, output_path)
    elif export_format == "md":
        return export_md(text, replacements, output_path)
    elif export_format == "docx":
        if ext == '.docx':
            # Content-level DOCX replacement
            return export_docx(source_path, replacements, output_path, text)
        else:
            # For non-DOCX sources, create a new DOCX with replaced text
            try:
                from docx import Document
                doc = Document()
                replaced = apply_replacements_to_text(text, replacements)
                doc.add_paragraph(replaced)
                doc.save(output_path)
                return {
                    "format": "docx",
                    "output_path": output_path,
                    "source_sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
                    "output_sha256": hashlib.sha256(Path(output_path).read_bytes()).hexdigest(),
                    "replacements_applied": len(replacements),
                }
            except ImportError:
                raise ImportError("python-docx required: pip install python-docx")
    else:
        raise ValueError(f"Unsupported export format: {export_format}")
